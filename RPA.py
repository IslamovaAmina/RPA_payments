import rpa as r
import win32com.client as win32
import clipboard as clp
import os
import docx
from docx2pdf import convert
from bs4 import BeautifulSoup
from PyPDF2 import PdfFileWriter, PdfFileReader
from zipfile import ZipFile

def login():
    r.url("url")
    if not r.read('login'):
        r.type('login')
    else:
        r.type('q', '[enter]')
    r.type('q', 'password')

def download_paysheet():
    r.click('HeaderTabs_Workers_Link')
    r.click('printActions')
    r.click('printPaysheets')
    r.dclick(300, 120)
    r.wait(5)

def find_file():
    while True:
        try:
            m_path = os.getcwd()
            m_path = m_path.split('\\')[1:3]
            user = m_path[1]
            m_path = os.path.join(*m_path)
            m_path = f"C:\{m_path}\Downloads"
            files = os.listdir(m_path)
            files = filter(lambda f: "Расчетный лист_" in f and f.split('.')[1] == 'docx', files)
            file = list(files)[0]
            return m_path, file, user
        except IndexError:
            continue


def move_to_buffer(m_path, file, user):
    start_path = m_path + "\\" + file
    dest_path = r"C:\Users" + f"\{user}\Gamma International Group\ONE STORY - Documents\Change" + \
                r"\Automation\Электроник\Buffer" + f"\{file}"
    print(start_path)
    print(dest_path)
    os.rename(start_path, dest_path)
    return start_path, dest_path


def mark_docx(dest_path):
    doc = docx.Document(dest_path)
    nums = []
    paragraphs = []
    order = []
    name_flag = 3
    for i, para in enumerate(doc.paragraphs):
        p = para.text
        if ("Общество с ограниченной " in p) and (i != 0):
            nums.append(i)
            name_flag = 3
        name_flag -= 1
        if name_flag == 0:
            order.append(p)
        # print(i, p)
        paragraphs.append(p)
    return nums, paragraphs, order, doc


def set_parts_to_new_page(nums, dest_path, doc):
    i = 0
    for num in nums:
        # print(doc.paragraphs[num].runs)
        # if not doc.paragraphs[num].runs:
        doc.paragraphs[num + i].insert_paragraph_before()
        doc.paragraphs[num + i].add_run()
        doc.paragraphs[num + i].runs[-1].add_break(docx.enum.text.WD_BREAK.PAGE)
        i += 1
        # doc.paragraphs[num].runs = [doc.paragraphs[num].runs[-1]] + doc.paragraphs[num].runs[:-1]
    # print(dest_path[:-1] + 'new.docx')
    raw_path = dest_path[:-4]
    new_doc = raw_path + 'new.docx'
    doc.save(new_doc)
    return raw_path, new_doc


def create_pdf(raw_path, docx_path, order, user):
    dest_path = raw_path + 'pdf'
    convert(docx_path, dest_path)
    with open(dest_path, "rb") as f:
        inputpdf = PdfFileReader(f)
        direct_path = r"C:\Users" + f"\{user}\Gamma International Group\ONE STORY - Documents\Change\Automation\Электроник\Buffer"
        for i in range(len(order)):
            output = PdfFileWriter()
            output.addPage(inputpdf.getPage(i))
            with open(direct_path + f"\{order[i]}.pdf", "wb") as outputStream:
                output.write(outputStream)
    return direct_path


def process_file(m_path, file, user):
    start_path, dest_path = move_to_buffer(m_path, file, user)
    nums, paragraphs, order, doc = mark_docx(dest_path)
    raw_path, new_doc = set_parts_to_new_page(nums, dest_path, doc)
    direct_path = create_pdf(raw_path, new_doc, order, user)
    return direct_path


def load_employees(user):
    employees_dest = r"C:\Users" + f"\{user}\Gamma International Group\ONE STORY - Documents\Change\Automation\Электроник\RPA Files\employees.csv"
    with open(employees_dest, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    lines[0] = lines[0][1:]
    lines = list(map(lambda x: x[1:-2].split(';'), lines))
    # print(lines)
    employee_dict = dict(lines)
    return employee_dict


def create_attachment(direct_path, file):
    target = direct_path + "\\" + file.split('.')[0]
    attachment = ZipFile(target + ".zip", 'w')
    file_path = target + ".pdf"
    attachment.write(file_path, os.path.basename(file_path))
    attachment.setpassword(b"123")
    attachment.close()
    return target + ".zip"


def send_emails(direct_path, employee_dict):
    outlook = win32.Dispatch('outlook.application')
    for file in os.listdir(direct_path):
        name = file.split('.')[0]
        if name in employee_dict.keys():
            file_path = create_attachment(direct_path, file)
            mail = outlook.CreateItem(0)
            # mail.To = 'islamova@onestory.pro"
            mail.To = employee_dict[name]
            mail.Subject = f'Расчетный лист'
            mail.Body = f'Уважаемый(-ая) {name},\n' + \
                        f'Направляю Вам расчетный листок по заработной плате.\nДокумент защищен паролем. ' + \
                        f'Для того, чтобы открыть документ, пожалуйста, наберите (пароль, номер паспорта).' + \
                        f'\nОбращаю внимание, что запросы на справки и ' + \
                        f'вопросы по начислениям вы можете направлять на e-mail: hr@onestory.pro.'
            mail.Attachments.Add(Source=file_path)
            mail.Send()


def clear_buffer():
    m_path = os.getcwd()
    m_path = m_path.split('\\')[1:3]
    user = m_path[1]
    direct_path = r"C:\Users" + f"\{user}\Gamma International Group\ONE STORY - Documents\Change\Automation\Электроник\Buffer"
    files = os.listdir(direct_path)
    for file in files:
        print(direct_path + "\\" + file)
        os.remove(direct_path + "\\" + file)


def main():
    m_path, file, user = find_file()
    direct_path = process_file(m_path, file, user)
    employee_dict = load_employees(user)
    send_emails(direct_path, employee_dict)
    clear_buffer()


if __name__ == "__main__":
    try:
        main()
    except:
        input()

